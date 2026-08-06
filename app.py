import streamlit as st
import requests
import io
from docx import Document

# Hàm hỗ trợ tạo file Word từ dữ liệu JSON
def create_word_doc(exam_data):
    doc = Document()
    
    doc.add_heading(exam_data.get("title", "Đề thi AI"), 0)
    doc.add_paragraph(f"Nguồn: {exam_data.get('source_material', 'Không xác định')} | Tổng điểm: {exam_data.get('total_points', 10)}")
    
    for q in exam_data.get("questions", []):
        doc.add_heading(f"Câu {q['id']}: {q['topic']}", level=1)
        doc.add_paragraph(f"Ngữ cảnh chung: {q['context']}")
        
        for sub_q in q.get("sub_questions", []):
            doc.add_heading(f"Câu {sub_q['id']} ({sub_q['points']} điểm) - Dạng: {sub_q['type']}", level=2)
            
            if sub_q['depends_on'] and str(sub_q['depends_on']).lower() != "none":
                doc.add_paragraph(f"(Sử dụng kết quả từ câu {sub_q['depends_on']})")
                
            doc.add_paragraph(f"Hỏi: {sub_q['prompt']}")
            doc.add_paragraph(f"Đáp án chi tiết:\n{sub_q['answer']}")
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# Cấu hình trang 
st.set_page_config(page_title="AI Exam Generator", layout="centered")

st.title("📝 Trợ lý tạo đề thi AI")
st.divider() 

st.header("Cấu hình Đề thi")

col1, col2 = st.columns([1.5, 1])

with col1:
    uploaded_file = st.file_uploader("Tải lên file bài giảng (PDF)", type=["pdf"])

with col2:
    topic = st.text_input("Chủ đề bài thi", placeholder="VD: Gradient Descent")
    num_questions = st.number_input("Số lượng câu hỏi lớn", min_value=1, max_value=5, value=3)

generate_btn = st.button("Tạo đề thi ngay", use_container_width=True)
st.divider()

# --- BƯỚC 1: XỬ LÝ GỌI API VÀ LƯU VÀO BỘ NHỚ ---
if generate_btn:
    if not uploaded_file or not topic:
        st.warning("Vui lòng tải lên file PDF và nhập chủ đề!")
    else:
        with st.spinner("AI đang đọc tài liệu và thiết kế đề thi... (khoảng 10-20 giây)"):
            try:
                files = {"file": (uploaded_file.name, uploaded_file.getvalue(), "application/pdf")}
                data = {"topic": topic, "num_questions": num_questions}
                
                response = requests.post("http://127.0.0.1:8000/generate-exam", files=files, data=data)
                
                if response.status_code == 200:
                    # LƯU KẾT QUẢ VÀO SESSION_STATE THAY VÌ IN RA LUÔN
                    st.session_state["exam_data"] = response.json()
                    st.session_state["current_topic"] = topic
                    st.success("Tạo đề thi thành công!")
                else:
                    st.error(f"Lỗi từ API: {response.text}")
                    
            except requests.exceptions.ConnectionError:
                st.error("Không thể kết nối đến Backend. Đảm bảo bạn đang chạy `uvicorn main:app` ở một terminal khác!")

# --- BƯỚC 2: HIỂN THỊ GIAO DIỆN ĐỘC LẬP VỚI NÚT BẤM ---
# Chỉ cần trong bộ nhớ có exam_data, khối này sẽ luôn được vẽ lại khi ấn "Hiện đáp án"
if "exam_data" in st.session_state:
    exam_data = st.session_state["exam_data"]
    saved_topic = st.session_state["current_topic"]
    
    st.header(exam_data.get("title", f"Đề thi: {saved_topic}"))
    st.caption(f"Nguồn: {exam_data.get('source_material', 'Không xác định')} | Tổng điểm: {exam_data.get('total_points', 10)}")
    st.divider()
    
    for q in exam_data.get("questions", []):
        st.subheader(f"Câu {q['id']}: {q['topic']}")
        st.info(f"**Ngữ cảnh chung:** {q['context']}")
        
        for sub_q in q.get("sub_questions", []):
            with st.expander(f"Câu {sub_q['id']} ({sub_q['points']} điểm) - Dạng: {sub_q['type']}"):
                if sub_q['depends_on'] and str(sub_q['depends_on']).lower() != "none":
                    st.markdown(f"*(Sử dụng kết quả từ câu **{sub_q['depends_on']}**)*")
                
                st.write(f"**Hỏi:** {sub_q['prompt']}")
                
                # THUẬT TOÁN NÚT BẤM HIỆN/ẨN ĐÁP ÁN
                state_key = f"ans_{q['id']}_{sub_q['id']}"
                
                if state_key not in st.session_state:
                    st.session_state[state_key] = False
                
                button_label = "Ẩn đáp án" if st.session_state[state_key] else "Hiện đáp án"
                
                if st.button(button_label, key=f"btn_{state_key}"):
                    st.session_state[state_key] = not st.session_state[state_key]
                    st.rerun() 
                
                if st.session_state[state_key]:
                    st.success(f"**Đáp án chi tiết:**\n{sub_q['answer']}")
    
    # NÚT TẢI FILE WORD NẰM TRONG KHỐI HIỂN THỊ
    st.divider()
    word_file = create_word_doc(exam_data)
    st.download_button(
        label="⬇️ Tải đề thi xuống (Word)",
        data=word_file,
        file_name=f"De_thi_{saved_topic.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )